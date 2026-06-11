"""Local document reader — extracts text from PDF, DOCX, HTML, and plain-text
formats on disk.

Provider-agnostic by design: the tool operates on local file paths, so it
pairs naturally with the mail integrations (``gmail_get_attachment``,
``proton_get_attachment``) but works on any file regardless of origin —
downloaded webpages, generated reports, files placed manually, etc.

Architecture notes:

* **Graceful degradation.** ``pypdf`` and ``python-docx`` are imported with
  try/except guards. If a library is missing, ``read_document`` returns a
  clear error for that specific format only — other formats still work. Same
  pattern as ``_HAS_GOOGLE`` / ``_HAS_MCP`` / ``_HAS_PROTONMAIL`` in the rest
  of the codebase.

* **Always-on tool.** Unlike Desktop/Browser/Meta/MCP/Google/Proton tools
  which are gated behind per-instruction checkboxes, ``read_document`` is in
  the always-available ``TOOLS`` list — sibling of ``csv_search``,
  ``run_command``, ``web_search``, ``fetch_webpage``. Any instruction that
  needs to read a local file gets it for free.

* **Format detection by extension first, content second.** File extensions
  are the primary signal because they encode the producer's intent. For
  files without recognisable extensions, we attempt UTF-8 decoding first
  (catches mislabelled text files) before falling back to a hex preview of
  the first 256 bytes.

* **Long-tail formats deliberately not supported.** XLSX, ZIP, audio, video,
  RTF, EPUB, etc. each have their own library. Adding them here would bloat
  the mixin and the dependency footprint. The tool description directs the
  agent to ``run_command`` with the appropriate CLI tool (unzip, pandoc,
  ffprobe, etc.) for those cases — keeping the mixin focused on the most
  common attachment formats.
"""

import json
import mimetypes
import os

from myagent.helpers import extract_text_from_html

_HAS_PYPDF = True
try:
    import pypdf
except Exception:
    _HAS_PYPDF = False

_HAS_DOCX = True
try:
    import docx  # python-docx package, imports as "docx"
except Exception:
    _HAS_DOCX = False


# Default text-extraction cap. Same value as gmail_read / proton_read body
# truncation so the agent has consistent ergonomics across "fetch a message"
# and "read its attachment" calls.
DEFAULT_MAX_CHARS = 50_000

# Extensions we read natively as plain text (no external library needed).
# Anything UTF-8-decodable will also fall through to this path even without
# a matching extension — see do_read_document's "try as text first" branch.
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".log", ".rst",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".env",
    ".csv", ".tsv", ".xml",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".sh", ".bash", ".zsh",
    ".bat", ".ps1", ".sql", ".rb", ".go", ".rs", ".java", ".kt",
    ".c", ".cpp", ".h", ".hpp", ".cs", ".swift", ".php",
}
HTML_EXTENSIONS = {".html", ".htm", ".xhtml"}


class DocumentMixin:
    """Local document reading tools.

    Single tool surface: ``do_read_document``. Provider-agnostic — operates
    on any local file path, regardless of source."""

    def do_read_document(self, params):
        path = params.get("path", "")
        max_chars = int(params.get("max_chars") or DEFAULT_MAX_CHARS)
        pages_str = params.get("pages")

        if not path:
            return "error: 'path' is required"
        if not os.path.exists(path):
            return f"error: file not found: {path}"
        if not os.path.isfile(path):
            return f"error: not a regular file: {path}"

        size = os.path.getsize(path)
        ext = os.path.splitext(path)[1].lower()
        mime_type, _ = mimetypes.guess_type(path)

        result = {
            "path": path,
            "size_bytes": size,
            "mime_type": mime_type or "application/octet-stream",
            "format": "unknown",
            "text": "",
            "text_truncated": False,
        }

        try:
            if ext == ".pdf":
                if not _HAS_PYPDF:
                    return (
                        "error: PDF support requires pypdf — install with "
                        "'pip install pypdf', then retry. Or fall back to "
                        "run_command with 'pdftotext <path> -' if poppler is installed."
                    )
                result["format"] = "pdf"
                self._read_pdf(path, result, max_chars, pages_str)
            elif ext == ".docx":
                if not _HAS_DOCX:
                    return (
                        "error: DOCX support requires python-docx — install with "
                        "'pip install python-docx', then retry. Or fall back to "
                        "run_command with 'pandoc <path> -t plain' if pandoc is installed."
                    )
                result["format"] = "docx"
                self._read_docx(path, result, max_chars)
            elif ext in HTML_EXTENSIONS:
                result["format"] = "html"
                with open(path, encoding="utf-8", errors="replace") as f:
                    html = f.read()
                text = extract_text_from_html(html)
                result["text"] = text[:max_chars]
                result["text_truncated"] = len(text) > max_chars
            elif ext in TEXT_EXTENSIONS:
                result["format"] = "text"
                with open(path, encoding="utf-8", errors="replace") as f:
                    text = f.read()
                result["text"] = text[:max_chars]
                result["text_truncated"] = len(text) > max_chars
            else:
                # Try as text first (catches mislabelled or extensionless text
                # files), fall back to a hex preview that lets the model at
                # least identify the file's nature.
                try:
                    with open(path, encoding="utf-8") as f:
                        text = f.read()
                    result["format"] = "text (assumed from content)"
                    result["text"] = text[:max_chars]
                    result["text_truncated"] = len(text) > max_chars
                except (UnicodeDecodeError, IsADirectoryError):
                    with open(path, "rb") as f:
                        head = f.read(256)
                    result["format"] = "binary"
                    result["text"] = (
                        f"Binary file ({size:,} bytes, MIME: {result['mime_type']}). "
                        f"Use run_command for format-specific tools (file, unzip, "
                        f"pandoc, ffprobe, etc.). First 256 bytes hex preview:\n"
                        f"{head.hex()}"
                    )
                    result["text_truncated"] = size > 256
                    result["binary_preview_hex"] = head.hex()
        except Exception as e:
            return f"error reading {path}: {type(e).__name__}: {e}"

        return json.dumps(result, indent=2, ensure_ascii=False)

    # ── Format-specific readers ─────────────────────────────────────────────

    def _read_pdf(self, path, result, max_chars, pages_str):
        """Extract text from a PDF using pypdf. Populates result in place.

        Per-page extraction is independent so a malformed page doesn't break
        the whole document — failed pages get a clearly-marked placeholder
        and the rest continues. Metadata is read from the document
        information dictionary if present."""
        reader = pypdf.PdfReader(path)
        result["page_count"] = len(reader.pages)

        # Page range parsing — defaults to all pages if not specified.
        if pages_str:
            pages = self._parse_pages_range(pages_str, len(reader.pages))
            if not pages:
                result["text"] = f"error: no valid pages in range {pages_str!r} (PDF has {len(reader.pages)} pages)"
                return
        else:
            pages = list(range(len(reader.pages)))

        # Metadata extraction. pypdf's PdfReader.metadata is None for PDFs
        # without an info dictionary; otherwise it's a dict-like object with
        # entries like /Title, /Author, /Creator, /Subject, /CreationDate.
        if reader.metadata:
            meta = {}
            for label, key in (
                ("title", "title"), ("author", "author"),
                ("creator", "creator"), ("subject", "subject"),
                ("creation_date", "creation_date"),
                ("modification_date", "modification_date"),
                ("producer", "producer"),
            ):
                try:
                    val = getattr(reader.metadata, key, None)
                except Exception:
                    val = None
                if val:
                    meta[label] = str(val)
            if meta:
                result["metadata"] = meta

        # Encryption check — if the PDF is password-protected, extract_text
        # would return empty for each page. Surface this explicitly rather
        # than silently returning a blank result.
        try:
            if reader.is_encrypted:
                # Attempt empty-password decryption (common for "owner
                # password only" PDFs that allow reading without auth).
                try:
                    reader.decrypt("")
                except Exception:
                    pass
                if reader.is_encrypted:
                    result["text"] = (
                        f"error: PDF is encrypted ({len(reader.pages)} pages). "
                        f"pypdf cannot decrypt without the password. If you have "
                        f"it, use run_command with 'qpdf --password=<pw> --decrypt "
                        f"<input> <output>' then read_document the decrypted output."
                    )
                    return
        except Exception:
            pass

        # Per-page extraction with per-page error isolation.
        parts = []
        for page_idx in pages:
            if page_idx < 0 or page_idx >= len(reader.pages):
                continue
            try:
                page_text = reader.pages[page_idx].extract_text() or ""
                parts.append(page_text)
            except Exception as e:
                parts.append(f"[page {page_idx + 1}: extraction error — {type(e).__name__}]")
        text = "\n\n".join(parts)
        result["text"] = text[:max_chars]
        result["text_truncated"] = len(text) > max_chars
        result["pages_extracted"] = len(pages)

    def _read_docx(self, path, result, max_chars):
        """Extract text from a .docx file using python-docx. Captures
        paragraphs AND table cells in document order. Tables are rendered as
        pipe-separated rows so column structure is preserved in plain text."""
        doc = docx.Document(path)

        # Capture paragraphs in document order, skipping empty ones.
        paragraph_count = 0
        parts = []
        for p in doc.paragraphs:
            paragraph_count += 1
            if p.text.strip():
                parts.append(p.text)

        # Tables — rendered as " | "-separated rows. python-docx exposes
        # tables via doc.tables (top-level) and via section iteration; for
        # this use case, top-level tables cover virtually all real documents.
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(row_cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

        result["paragraph_count"] = paragraph_count
        result["table_count"] = table_count

        # Metadata via core properties (title, author, etc.)
        try:
            cp = doc.core_properties
            meta = {}
            for attr in ("title", "author", "subject", "created", "modified", "last_modified_by", "revision"):
                val = getattr(cp, attr, None)
                if val:
                    meta[attr] = str(val)
            if meta:
                result["metadata"] = meta
        except Exception:
            pass

        text = "\n".join(parts)
        result["text"] = text[:max_chars]
        result["text_truncated"] = len(text) > max_chars

    # ── Helpers ─────────────────────────────────────────────────────────────

    @staticmethod
    def _parse_pages_range(spec, total):
        """Parse a 1-indexed page-range string like '1-5', '3', or '1,3,5-7'
        into a list of 0-indexed page numbers. Out-of-range values are
        silently dropped so the caller gets at least a partial result rather
        than an error. Returns [] for completely unparseable input."""
        result = []
        for part in (spec or "").split(","):
            part = part.strip()
            if not part:
                continue
            if "-" in part:
                start_s, _, end_s = part.partition("-")
                try:
                    start_n = int(start_s.strip()) - 1
                    end_n = int(end_s.strip()) - 1
                except ValueError:
                    continue
                result.extend(i for i in range(start_n, end_n + 1) if 0 <= i < total)
            else:
                try:
                    n = int(part) - 1
                    if 0 <= n < total:
                        result.append(n)
                except ValueError:
                    continue
        # Deduplicate while preserving order
        seen = set()
        deduped = []
        for n in result:
            if n not in seen:
                seen.add(n)
                deduped.append(n)
        return deduped
